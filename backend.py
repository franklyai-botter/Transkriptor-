"""
NeuralNautic Transcriptor - Backend
FastAPI Server: Transkription, Stille-Erkennung, Folien-Extraktion, Vision-Analyse
"""

import os
import json
import uuid
import shutil
import subprocess
import tempfile
import base64
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent / ".env")

import imageio_ffmpeg
import cv2

# FFmpeg Pfad (imageio-ffmpeg Bundle — kein separates Install nötig)
FFMPEG = imageio_ffmpeg.get_ffmpeg_exe()

# FFmpeg-Verzeichnis in PATH eintragen, damit Whisper es findet
_ffmpeg_dir = str(Path(FFMPEG).parent)
if _ffmpeg_dir not in os.environ.get("PATH", ""):
    os.environ["PATH"] = _ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")

import numpy as np
from faster_whisper import WhisperModel
import anthropic
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Inches
import markdown

app = FastAPI(title="NeuralNautic Transcriptor")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Arbeitsverzeichnis
WORK_DIR = Path("transcriptor_jobs")
WORK_DIR.mkdir(exist_ok=True)

# Jobs status speichern
jobs: dict = {}

# Whisper Modelle via faster-whisper (CTranslate2) — ~3-4x schneller als openai-whisper
# auf CPU bei gleicher Qualitaet. Es bleibt immer nur EIN Modell im RAM; beim Wechsel
# wird neu geladen (nach dem ersten Download aus dem Disk-Cache, schnell).
ALLOWED_WHISPER_MODELS = ("tiny", "base", "small", "medium", "large")
DEFAULT_WHISPER_MODEL = "large"
# Mapping auf die faster-whisper-Repo-IDs; "large" -> large-v3 (beste Qualitaet)
_FW_MODEL_ID = {
    "tiny": "tiny", "base": "base", "small": "small",
    "medium": "medium", "large": "large-v3",
}
_whisper_cache = {"name": None, "model": None}

def get_whisper_model(model_name: str = DEFAULT_WHISPER_MODEL):
    if model_name not in ALLOWED_WHISPER_MODELS:
        model_name = DEFAULT_WHISPER_MODEL
    if _whisper_cache["name"] != model_name:
        fw_id = _FW_MODEL_ID[model_name]
        print(f"[Whisper] Lade faster-whisper Modell: {fw_id} (erster Download kann dauern)", flush=True)
        # int8 = schnell + speichersparend auf CPU
        _whisper_cache["model"] = WhisperModel(fw_id, device="cpu", compute_type="int8")
        _whisper_cache["name"] = model_name
        print(f"[Whisper] Modell geladen: {fw_id}", flush=True)
    return _whisper_cache["model"]


# ── Audio Extraktion ──────────────────────────────────────────────────────────

def extract_audio(video_path: Path, out_path: Path) -> bool:
    """Extrahiert Audio aus Video mit FFmpeg"""
    result = subprocess.run([
        FFMPEG, "-y", "-i", str(video_path),
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        str(out_path)
    ], capture_output=True, text=True)
    return result.returncode == 0


# ── Stille-Erkennung ─────────────────────────────────────────────────────────

def detect_silence(audio_path: Path, min_silence_sec: float = 2.0) -> list[dict]:
    """Findet Stille-Segmente via FFmpeg silencedetect"""
    result = subprocess.run([
        FFMPEG, "-i", str(audio_path),
        "-af", f"silencedetect=noise=-40dB:d={min_silence_sec}",
        "-f", "null", "-"
    ], capture_output=True, text=True)

    output = result.stderr
    silent_segments = []
    starts = []

    for line in output.split('\n'):
        if 'silence_start' in line:
            try:
                t = float(line.split('silence_start:')[1].strip())
                starts.append(t)
            except:
                pass
        elif 'silence_end' in line:
            try:
                parts = line.split('silence_end:')[1].strip().split('|')
                end = float(parts[0].strip())
                if starts:
                    start = starts.pop(0)
                    silent_segments.append({
                        "start": round(start, 2),
                        "end": round(end, 2),
                        "duration": round(end - start, 2)
                    })
            except:
                pass

    return silent_segments


# ── Folien-Extraktion ─────────────────────────────────────────────────────────

def extract_slides(video_path: Path, out_dir: Path, threshold: float = 0.93,
                   progress_cb=None) -> list[dict]:
    """Erkennt Szenenaenderungen und extrahiert Folien/Keyframes.

    progress_cb(progress_0_to_1, slides_found_so_far): Live-Update.
    Skip-Logik: nach 20% der Video-Laenge ohne einen einzigen Folien-Wechsel
    wird der Scan abgebrochen und [] zurueckgegeben (typisches Webcam/Sprech-Video).
    """
    out_dir.mkdir(exist_ok=True)
    cap = cv2.VideoCapture(str(video_path))
    fps          = cap.get(cv2.CAP_PROP_FPS) or 30.0
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    slides = []
    prev_frame      = None
    frame_idx       = 0
    slide_count     = 0
    last_slide_time = -3.0
    last_pct_emit   = -1
    early_check_at  = int(total_frames * 0.20) if total_frames > 0 else 99999

    while True:
        ret, frame = cap.read()
        if not ret:
            break

        # Nur jeden 15. Frame pruefen (Performance + weniger False Positives)
        if frame_idx % 15 == 0:
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            gray = cv2.resize(gray, (160, 90))

            if prev_frame is not None:
                diff      = cv2.absdiff(prev_frame, gray)
                score     = 1.0 - (diff.mean() / 255.0)
                timestamp = frame_idx / fps
                if score < threshold and (timestamp - last_slide_time) >= 3.0:
                    filename = f"slide_{slide_count:04d}_{int(timestamp)}s.jpg"
                    cv2.imwrite(str(out_dir / filename), frame)
                    slides.append({
                        "timestamp": round(timestamp, 2),
                        "timestamp_fmt": _fmt_time(timestamp),
                        "file": filename,
                        "frame": frame_idx
                    })
                    slide_count    += 1
                    last_slide_time = timestamp

            prev_frame = gray

            # Live-Progress alle 2 % melden
            if progress_cb and total_frames > 0:
                pct = int(frame_idx * 100 / total_frames)
                if pct - last_pct_emit >= 2:
                    progress_cb(frame_idx / total_frames, slide_count)
                    last_pct_emit = pct

        # Early-Skip: nach 20 % der Video-Laenge keine einzige Folie -> Abbruch
        if frame_idx >= early_check_at and slide_count == 0:
            cap.release()
            return []

        frame_idx += 1

    cap.release()

    # Ersten Frame nur hinzufuegen wenn ueberhaupt Folien gefunden wurden
    if slide_count > 0:
        cap2 = cv2.VideoCapture(str(video_path))
        ret, first_frame = cap2.read()
        if ret:
            filename = "slide_first_0s.jpg"
            cv2.imwrite(str(out_dir / filename), first_frame)
            slides.insert(0, {"timestamp": 0.0, "timestamp_fmt": "00:00",
                              "file": filename, "frame": 0})
        cap2.release()

    return slides


def _fmt_time(seconds: float) -> str:
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m:02d}:{s:02d}"


# ── Vision-Analyse ────────────────────────────────────────────────────────────

class VisionAuthError(Exception):
    """API-Key ungueltig — keine weiteren Vision-Calls in diesem Job versuchen."""


def analyze_frame(image_path: Path) -> dict:
    """Sendet einen Frame an Claude Vision — erkennt Typ und extrahiert Text.

    Wirft `VisionAuthError` bei 401/403, damit der aufrufende Job-Loop
    sofort abbricht statt jeden weiteren Frame zu probieren.
    Andere Fehler werden als {"type":"other","content":""} verschluckt.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"type": "other", "content": ""}

    try:
        with open(image_path, "rb") as f:
            img_b64 = base64.standard_b64encode(f.read()).decode("utf-8")
    except Exception:
        return {"type": "other", "content": ""}

    client = anthropic.Anthropic(api_key=api_key)
    try:
        message = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=1024,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}
                    },
                    {
                        "type": "text",
                        "text": (
                            "Analysiere dieses Bild aus einem Screen-Recording.\n\n"
                            "1. Bestimme den Typ: 'mural' (digitales Whiteboard mit Sticky Notes, Karten, Clustern), "
                            "'slide' (Praesentationsfolie), oder 'other' (alles andere).\n\n"
                            "2. Wenn 'mural' oder 'slide': Extrahiere den gesamten sichtbaren Text strukturiert.\n"
                            "   - Bei Mural: Cluster/Bereiche als Ueberschriften, Sticky Notes als Bullet Points\n"
                            "   - Bei Slide: Titel fett, dann Bullet Points\n\n"
                            "Antworte NUR in diesem JSON-Format (kein Markdown drumherum):\n"
                            '{"type": "mural|slide|other", "content": "extrahierter Text oder leer"}'
                        )
                    }
                ]
            }]
        )
        print(f"[Vision] Frame analysiert mit Modell: {message.model}", flush=True)
    except anthropic.AuthenticationError:
        # API-Key ungueltig -> Job-Loop signalisieren dass Vision unbrauchbar ist
        raise VisionAuthError("Anthropic API-Key ungueltig (401)")
    except Exception:
        # Netzwerk, Rate-Limit, Timeout etc. -> nur diesen Frame skippen
        return {"type": "other", "content": ""}

    try:
        raw = message.content[0].text.strip()
        return json.loads(raw)
    except Exception:
        return {"type": "other", "content": ""}


# ── Transkription ─────────────────────────────────────────────────────────────

def transcribe_audio(audio_path: Path, model_name: str = DEFAULT_WHISPER_MODEL) -> dict:
    """Transkribiert Audio mit faster-whisper. Gibt ein openai-whisper-kompatibles Dict
    zurueck (segments/language/text), damit build_transcript_data unveraendert bleibt.
    faster-whisper dekodiert die WAV selbst (gebuendeltes PyAV) — kein ffmpeg-Aufruf noetig."""
    model = get_whisper_model(model_name)

    segments_gen, info = model.transcribe(str(audio_path), language="de")

    segments = []
    texts = []
    for seg in segments_gen:  # Generator — hier laeuft die eigentliche Transkription
        segments.append({"start": seg.start, "end": seg.end, "text": seg.text})
        texts.append(seg.text)

    return {
        "segments": segments,
        "language": info.language,
        "text": "".join(texts),
    }


# ── Export-Funktionen ─────────────────────────────────────────────────────────

def build_transcript_data(whisper_result: dict, silence: list, slides: list) -> dict:
    """Kombiniert alle Daten zu einem strukturierten Transkript"""
    segments = []
    for seg in whisper_result.get("segments", []):
        segments.append({
            "start": round(seg["start"], 2),
            "end": round(seg["end"], 2),
            "start_fmt": _fmt_time(seg["start"]),
            "end_fmt": _fmt_time(seg["end"]),
            "text": seg["text"].strip(),
            "type": "speech"
        })

    # Stille-Segmente einmischen
    for s in silence:
        if s["duration"] >= 2.0:
            segments.append({
                "start": s["start"],
                "end": s["end"],
                "start_fmt": _fmt_time(s["start"]),
                "end_fmt": _fmt_time(s["end"]),
                "text": f"[Kein Ton – {s['duration']:.0f}s Stille]",
                "type": "silence"
            })

    segments.sort(key=lambda x: x["start"])

    # Zugehörige Folie für jedes Segment finden
    for seg in segments:
        seg["slide"] = None
        seg["visual_type"] = None
        seg["visual_content"] = None
        for slide in reversed(slides):
            if slide["timestamp"] <= seg["start"]:
                seg["slide"] = slide["file"]
                seg["visual_type"] = slide.get("visual_type")
                seg["visual_content"] = slide.get("visual_content")
                break

    return {
        "language": whisper_result.get("language", "unknown"),
        "segments": segments,
        "slides": slides,
        "silence": silence,
        "full_text": whisper_result.get("text", "").strip()
    }


def export_markdown(data: dict, out_path: Path):
    lines = [
        f"# Transkript",
        f"",
        f"**Sprache:** {data['language']}",
        f"",
        "---",
        "",
    ]

    last_visual_content = None  # vermeidet Duplikate wenn mehrere Segmente dieselbe Folie haben

    for seg in data["segments"]:
        # Visuellen Inhalt einmalig einfügen wenn er sich ändert
        visual_content = seg.get("visual_content")
        visual_type = seg.get("visual_type")
        if visual_content and visual_content != last_visual_content and visual_type in ("mural", "slide"):
            label = "Mural-Board" if visual_type == "mural" else "Folie"
            lines.append(f"### {label} @ {seg['start_fmt']}")
            lines.append("")
            lines.append(visual_content)
            lines.append("")
            last_visual_content = visual_content

        if seg["type"] == "silence":
            lines.append(f"> *[{seg['start_fmt']} – {seg['end_fmt']}] {seg['text']}*")
            lines.append("")
        else:
            lines.append(f"**[{seg['start_fmt']}]** {seg['text']}")
            lines.append("")

    out_path.write_text("\n".join(lines), encoding="utf-8")


def export_html(data: dict, out_path: Path, slides_dir: Path):
    segments_html = ""
    for seg in data["segments"]:
        slide_html = ""
        if seg.get("slide"):
            slide_html = f'<div class="seg-slide"><img src="slides/{seg["slide"]}" class="slide-thumb" alt="Folie"></div>'

        if seg["type"] == "silence":
            segments_html += f"""
            <div class="segment silence">
              <span class="seg-time">{seg['start_fmt']} – {seg['end_fmt']}</span>
              <div class="seg-content"><span class="seg-text">~ {seg['text']}</span></div>
            </div>"""
        else:
            segments_html += f"""
            <div class="segment speech">
              <span class="seg-time">[{seg['start_fmt']}]</span>
              <div class="seg-content"><span class="seg-text">{seg['text']}</span>{slide_html}</div>
            </div>"""

    html = f"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<title>NeuralNautic Transkript</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Italiana&family=Inter:wght@300;400;500;600&family=JetBrains+Mono:wght@400;500&display=swap">
<style>
  :root {{
    --ink-abyss:    #061419;
    --ink-deep:     #0A2028;
    --ink-tide:     #0E2B36;
    --glow-faint:   #2D8A9A;
    --glow-cyan:    #3FD4E0;
    --glow-aqua:    #6FEFF5;
    --silver-mid:   #9C9CA4;
    --silver-bright:#CDCED2;
    --silver-white: #F5F5F2;
    --signal-alert: #F08A7A;
    --border-1:     rgba(205,206,210,0.10);
    --border-2:     rgba(205,206,210,0.18);
    --font-display: 'Italiana', serif;
    --font-body:    'Inter', system-ui, sans-serif;
    --font-mono:    'JetBrains Mono', monospace;
  }}
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{
    margin: 0;
    background: var(--ink-deep);
    color: var(--silver-white);
    font-family: var(--font-body);
    font-size: 16px;
    line-height: 1.5;
    padding: 2.5rem 2rem;
    -webkit-font-smoothing: antialiased;
  }}
  .nn-container {{ max-width: 960px; margin: 0 auto; }}
  .nn-header {{
    border-bottom: 1px solid var(--border-1);
    padding-bottom: 1.4rem;
    margin-bottom: 2rem;
  }}
  h1 {{
    font-family: var(--font-display);
    font-weight: 400;
    font-size: 2.2rem;
    color: var(--silver-white);
    letter-spacing: 0.04em;
    margin: 0;
  }}
  .nn-sub {{
    font-family: var(--font-mono);
    font-size: 0.72rem;
    letter-spacing: 0.18em;
    text-transform: uppercase;
    color: var(--silver-mid);
    margin-top: 0.4rem;
  }}
  .nn-meta {{
    font-family: var(--font-mono);
    font-size: 0.78rem;
    color: var(--silver-bright);
    letter-spacing: 0.05em;
    margin-top: 1rem;
  }}
  .nn-meta strong {{ color: var(--glow-cyan); font-weight: 500; }}
  .segment {{
    display: flex;
    gap: 1rem;
    padding: 0.9rem 1rem;
    margin-bottom: 0.5rem;
    border-radius: 8px;
    align-items: flex-start;
  }}
  .segment.speech {{
    border-left: 3px solid var(--glow-faint);
    background: rgba(14,43,54,0.55);
  }}
  .segment.silence {{
    border-left: 3px solid var(--signal-alert);
    background: rgba(240,138,122,0.06);
  }}
  .seg-time {{
    font-family: var(--font-mono);
    color: var(--glow-faint);
    font-size: 0.74rem;
    letter-spacing: 0.05em;
    white-space: nowrap;
    min-width: 90px;
    font-variant-numeric: tabular-nums;
    padding-top: 2px;
  }}
  .seg-content {{ flex: 1; }}
  .seg-text {{ font-size: 0.95rem; line-height: 1.6; }}
  .segment.silence .seg-text {{
    color: var(--signal-alert);
    font-family: var(--font-mono);
    font-size: 0.82rem;
    letter-spacing: 0.04em;
  }}
  .seg-slide {{ margin-top: 0.6rem; }}
  .slide-thumb {{
    max-width: 240px;
    border-radius: 8px;
    border: 1px solid var(--border-2);
  }}
  .nn-footer {{
    margin-top: 2.5rem;
    padding-top: 1.2rem;
    border-top: 1px solid var(--border-1);
    font-family: var(--font-mono);
    font-size: 0.68rem;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: var(--silver-mid);
    text-align: center;
  }}
</style>
</head>
<body>
<div class="nn-container">
  <header class="nn-header">
    <h1>NeuralNautic Transkript</h1>
    <div class="nn-sub">Transcriptor · Whisper Pipeline</div>
    <div class="nn-meta">Sprache: <strong>{data['language']}</strong></div>
  </header>
  {segments_html}
  <footer class="nn-footer">NeuralNautic — Lokale Transkription</footer>
</div>
</body>
</html>"""
    out_path.write_text(html, encoding="utf-8")


def export_pdf(data: dict, out_path: Path, slides_dir: Path):
    from reportlab.lib.colors import HexColor

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    # NN-Brand colors
    cyan = HexColor("#3FD4E0")
    cyan_dark = HexColor("#2D8A9A")
    silver_mid = HexColor("#9C9CA4")
    alert = HexColor("#F08A7A")
    ink_deep = HexColor("#0A2028")

    # Header: Logo-Star + Title + Sub
    logo_path = Path(__file__).parent / "frontend" / "logo-star.png"
    if logo_path.exists():
        try:
            story.append(RLImage(str(logo_path), width=1.4 * cm, height=1.4 * cm))
            story.append(Spacer(1, 0.2 * cm))
        except Exception:
            pass

    title_style = ParagraphStyle(
        'nn-title', parent=styles['Heading1'],
        fontName='Helvetica', fontSize=20, leading=24,
        textColor=ink_deep, spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        'nn-sub', parent=styles['Normal'],
        fontName='Courier', fontSize=8, leading=10,
        textColor=silver_mid, spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        'nn-meta', parent=styles['Normal'],
        fontName='Courier', fontSize=8, leading=10,
        textColor=cyan_dark, spaceAfter=14,
    )
    speech_style = ParagraphStyle(
        'nn-speech', parent=styles['Normal'],
        fontName='Helvetica', fontSize=10, leading=14,
        spaceAfter=4,
    )
    silence_style = ParagraphStyle(
        'nn-silence', parent=styles['Italic'],
        fontName='Courier-Oblique', fontSize=9, leading=12,
        textColor=alert, spaceAfter=4,
    )
    ts_style = ParagraphStyle(
        'nn-ts', parent=styles['Normal'],
        fontName='Courier-Bold', fontSize=8, leading=10,
        textColor=cyan_dark, spaceAfter=2,
    )

    story.append(Paragraph("NeuralNautic Transkript", title_style))
    story.append(Paragraph("TRANSCRIPTOR &middot; WHISPER PIPELINE", sub_style))
    story.append(Paragraph(f"SPRACHE: {data['language'].upper()}", meta_style))

    for seg in data["segments"]:
        if seg["type"] == "silence":
            story.append(Paragraph(
                f"[{seg['start_fmt']}&ndash;{seg['end_fmt']}] ~ {seg['text']}",
                silence_style,
            ))
        else:
            story.append(Paragraph(f"[{seg['start_fmt']}]", ts_style))
            story.append(Paragraph(seg["text"], speech_style))

        if seg.get("slide"):
            slide_path = slides_dir / seg["slide"]
            if slide_path.exists():
                try:
                    story.append(RLImage(str(slide_path), width=8 * cm, height=4.5 * cm))
                except Exception:
                    pass
        story.append(Spacer(1, 0.2 * cm))

    doc.build(story)


def export_docx(data: dict, out_path: Path, slides_dir: Path):
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Header: Logo-Star + Title
    logo_path = Path(__file__).parent / "frontend" / "logo-star.png"
    if logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Inches(0.5))
        except Exception:
            pass

    heading = doc.add_heading("NeuralNautic Transkript", 0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x20, 0x28)

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run("TRANSCRIPTOR · WHISPER PIPELINE")
    sub_run.font.name = "Consolas"
    sub_run.font.size = Pt(8)
    sub_run.font.color.rgb = RGBColor(0x9C, 0x9C, 0xA4)

    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"SPRACHE: {data['language'].upper()}")
    meta_run.font.name = "Consolas"
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(0x2D, 0x8A, 0x9A)

    for seg in data["segments"]:
        if seg["type"] == "silence":
            p = doc.add_paragraph()
            ts = p.add_run(f"[{seg['start_fmt']}–{seg['end_fmt']}] ")
            ts.bold = True
            ts.font.name = "Consolas"
            ts.font.size = Pt(9)
            ts.font.color.rgb = RGBColor(0xF0, 0x8A, 0x7A)
            txt = p.add_run(f"~ {seg['text']}")
            txt.italic = True
            txt.font.color.rgb = RGBColor(0xF0, 0x8A, 0x7A)
        else:
            p = doc.add_paragraph()
            ts = p.add_run(f"[{seg['start_fmt']}] ")
            ts.bold = True
            ts.font.name = "Consolas"
            ts.font.size = Pt(9)
            ts.font.color.rgb = RGBColor(0x2D, 0x8A, 0x9A)
            p.add_run(seg["text"])

        if seg.get("slide"):
            slide_path = slides_dir / seg["slide"]
            if slide_path.exists():
                try:
                    doc.add_picture(str(slide_path), width=Inches(4))
                except Exception:
                    pass

    doc.save(str(out_path))


# ── Background Job ─────────────────────────────────────────────────────────────

def run_job(job_id: str, video_path: Path):
    job_dir = WORK_DIR / job_id
    job_dir.mkdir(exist_ok=True)
    slides_dir = job_dir / "slides"
    mode = jobs.get(job_id, {}).get("mode", "full")  # 'full' | 'text_only'
    whisper_model = jobs.get(job_id, {}).get("whisper_model", DEFAULT_WHISPER_MODEL)

    def cancelled() -> bool:
        return jobs.get(job_id, {}).get("status") == "cancelled"

    try:
        def update(step: str, progress: int, msg: str = ""):
            if cancelled():
                raise RuntimeError("Job abgebrochen")
            jobs[job_id]["step"] = step
            jobs[job_id]["progress"] = progress
            jobs[job_id]["message"] = msg

        # ── Q1 (0-25%): Audio + Stille ──────────────────────────────────
        update("audio", 3, "Extrahiere Audio...")
        audio_path = job_dir / "audio.wav"
        if not extract_audio(video_path, audio_path):
            raise RuntimeError("FFmpeg Audio-Extraktion fehlgeschlagen")

        update("audio", 15, "Erkenne Stille-Segmente...")
        silence = detect_silence(audio_path)
        update("audio", 25, f"Audio analysiert ({len(silence)} Stille-Segmente)")

        # ── Q2 (25-50%): Folien + Vision — nur im 'full'-Mode ───────────
        if mode == "full":
            update("slides", 27, "Suche Folien-Wechsel...")

            def _slide_progress(p, found):
                pct = 27 + int(p * 18)  # 27..45
                update("slides", pct, f"Folien-Scan {int(p*100)}% — {found} gefunden")

            slides = extract_slides(video_path, slides_dir, progress_cb=_slide_progress)

            if not slides:
                update("slides", 50, "Keine Folien erkannt — uebersprungen.")
            elif os.environ.get("ANTHROPIC_API_KEY", ""):
                n = len(slides)
                vision_disabled = False
                for i, slide in enumerate(slides):
                    pct = 45 + int((i + 1) / n * 5)
                    img_path = slides_dir / slide["file"]
                    if vision_disabled or not img_path.exists():
                        slide["visual_type"]    = "other"
                        slide["visual_content"] = ""
                        continue
                    update("slides", pct, f"Vision-Analyse {i+1}/{n}")
                    try:
                        analysis = analyze_frame(img_path)
                        slide["visual_type"]    = analysis.get("type", "other")
                        slide["visual_content"] = analysis.get("content", "")
                    except VisionAuthError:
                        vision_disabled = True
                        slide["visual_type"]    = "other"
                        slide["visual_content"] = ""
                        update("slides", 50,
                               "Anthropic API-Key ungueltig — Vision uebersprungen.")
            else:
                for slide in slides:
                    slide["visual_type"]    = "other"
                    slide["visual_content"] = ""

            transcribe_start_pct = 52
            transcribe_end_pct   = 75
            export_start_pct     = 78
        else:
            # text_only: kein Folien-Scan, kein Vision; Whisper bekommt 25..90 %
            slides = []
            update("transcribe", 27, "Nur-Text-Modus — keine Folien-Analyse.")
            transcribe_start_pct = 27
            transcribe_end_pct   = 90
            export_start_pct     = 92

        # ── Q3 (Whisper) ────────────────────────────────────────────────
        whisper_msg = (
            f"Whisper ({whisper_model}) laeuft... ({len(slides)} Folien)"
            if slides else f"Whisper ({whisper_model}) laeuft..."
        )
        update("transcribe", transcribe_start_pct, whisper_msg)
        whisper_result = transcribe_audio(audio_path, whisper_model)
        update("transcribe", transcribe_end_pct, "Transkription abgeschlossen.")

        # ── Q4 (Export) ─────────────────────────────────────────────────
        update("export", export_start_pct, "Erstelle Transkript-Struktur...")
        transcript_data = build_transcript_data(whisper_result, silence, slides)

        update("export", 82, "Exportiere Markdown...")
        export_markdown(transcript_data, job_dir / "transcript.md")
        update("export", 87, "Exportiere HTML...")
        export_html(transcript_data, job_dir / "transcript.html", slides_dir)
        update("export", 92, "Exportiere PDF...")
        export_pdf(transcript_data, job_dir / "transcript.pdf", slides_dir)
        update("export", 97, "Exportiere DOCX...")
        export_docx(transcript_data, job_dir / "transcript.docx", slides_dir)

        # JSON für Frontend
        (job_dir / "transcript.json").write_text(
            json.dumps(transcript_data, ensure_ascii=False, indent=2), encoding="utf-8"
        )

        # ── Zwischenprodukt aufraeumen ───────────────────────────────────
        # audio.wav wird nur von detect_silence() und transcribe_audio() gelesen, beide
        # sind hier durch. Ohne dieses Loeschen bleibt je Job eine WAV von 80-370 MB
        # liegen: am 2026-08-02 waren es 39 Stueck / 7,2 GB, und der Ordner waechst mit
        # jedem Lauf weiter. Verlustfrei, weil extract_audio() sie am Anfang jedes Jobs
        # aus dem Video neu erzeugt.
        #
        # Bewusst NUR im Erfolgspfad: schlaegt ein Job fehl, bleibt die WAV fuer die
        # Fehlersuche liegen. Und bewusst in try/except — ein gesperrtes File (Virenscanner,
        # offener Player) darf einen fertigen Job nicht nachtraeglich auf "error" setzen.
        try:
            audio_path.unlink(missing_ok=True)
        except OSError as cleanup_err:
            print(f"[{job_id}] audio.wav nicht loeschbar: {cleanup_err}")

        jobs[job_id]["status"] = "done"
        jobs[job_id]["progress"] = 100
        jobs[job_id]["message"] = "Fertig!"
        jobs[job_id]["data"] = transcript_data

    except Exception as e:
        jobs[job_id]["status"] = "error"
        jobs[job_id]["message"] = str(e)


# ── API Routes ────────────────────────────────────────────────────────────────

@app.post("/upload")
async def upload_video(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    mode: str = Form("full"),
    whisper_model: str = Form(DEFAULT_WHISPER_MODEL),
):
    """mode = 'full' (Audio+Stille+Folien+Vision+Whisper+Export)
       mode = 'text_only' (Audio+Stille+Whisper+Export — Folien-Scan uebersprungen)
       whisper_model = tiny | base | small | medium | large"""
    if mode not in ("full", "text_only"):
        mode = "full"
    if whisper_model not in ALLOWED_WHISPER_MODELS:
        whisper_model = DEFAULT_WHISPER_MODEL

    job_id = str(uuid.uuid4())[:8]
    job_dir = WORK_DIR / job_id
    job_dir.mkdir(exist_ok=True)

    video_path = job_dir / file.filename
    with open(video_path, "wb") as f:
        content = await file.read()
        f.write(content)

    jobs[job_id] = {
        "id": job_id,
        "filename": file.filename,
        "mode": mode,
        "whisper_model": whisper_model,
        "status": "running",
        "step": "start",
        "progress": 0,
        "message": "Starte...",
        "data": None
    }

    background_tasks.add_task(run_job, job_id, video_path)
    return {"job_id": job_id, "mode": mode, "whisper_model": whisper_model}


@app.get("/status/{job_id}")
def get_status(job_id: str):
    if job_id not in jobs:
        return JSONResponse(status_code=404, content={"error": "Job nicht gefunden"})
    return jobs[job_id]


@app.get("/download/{job_id}/{format}")
def download(job_id: str, format: str):
    allowed = {"md", "html", "pdf", "docx", "json"}
    if format not in allowed:
        return JSONResponse(status_code=400, content={"error": "Ungültiges Format"})
    path = WORK_DIR / job_id / f"transcript.{format}"
    if not path.exists():
        return JSONResponse(status_code=404, content={"error": "Datei nicht gefunden"})
    return FileResponse(str(path), filename=f"transcript.{format}")


@app.get("/slides/{job_id}/{filename}")
def get_slide(job_id: str, filename: str):
    path = WORK_DIR / job_id / "slides" / filename
    if not path.exists():
        return JSONResponse(status_code=404, content={"error": "Folie nicht gefunden"})
    return FileResponse(str(path))


@app.post("/cancel/{job_id}")
def cancel_job(job_id: str):
    """Markiert Job als abgebrochen — der laufende run_job prüft das beim nächsten update() und steigt aus."""
    if job_id not in jobs:
        return JSONResponse(status_code=404, content={"error": "Job nicht gefunden"})
    if jobs[job_id]["status"] == "running":
        jobs[job_id]["status"] = "cancelled"
        jobs[job_id]["message"] = "Abgebrochen"
    return {"ok": True, "status": jobs[job_id]["status"]}


@app.get("/health")
def health():
    return {"ok": True}


@app.post("/shutdown")
def shutdown():
    """Eject-Button: stoppt den Server hart. start.bat kann ihn neu starten."""
    import threading, os, signal
    def _kill():
        import time
        time.sleep(0.3)
        os.kill(os.getpid(), signal.SIGTERM if hasattr(signal, "SIGTERM") else signal.SIGINT)
    threading.Thread(target=_kill, daemon=True).start()
    return {"ok": True, "message": "Server stoppt..."}


# Statische Dateien (Frontend)
app.mount("/", StaticFiles(directory="frontend", html=True), name="frontend")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=5678, reload=False)
